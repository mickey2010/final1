""" Final Project - This program will generate a recipe book of taco, including a random picture of taco

as well s three randomly generated recipes """



# import docx library to create and change a word document

import docx



# this requests library that  will get the url

import requests



# import image and font

from PIL import Image, ImageDraw, ImageFont



# open an image that will be use later in the project

Image = Image.open('photo-1552332386-f8dd00dc2f85.jpeg')



# resizing the image not more or less than 800

Image.thumbnail((800, 800))



# use the same font color for text and image

font = ImageFont.truetype('DejaVuSans.ttf', 20)



# download Dejavu from the website upon instructed



imgDraw = ImageDraw.Draw(Image)  # typing(draw) the text to the image

imgDraw.text([200, 80], 'Random Taco Cookbook', fill='purple', font=font)

# the ' Random text cookbook' text [200, 80] and fill purple in the text



Image.save('taco_thumbnails.jpeg')  # save the image



# save taco image and naming it

Image.save('modified_taco.jpg')



tacoBook = []  # list for creating three dictionaries of taco



# Loop for URL result and three tacos recipe

for i in range(3):

    tacoURL = requests.get('https://taco-1150.herokuapp.com/random/?full_taco=true').json()

    tacoBook.append(tacoURL)  # appending the tacoBook and URL

    print(tacoBook)  # printing





# creating a word document

tacoWord = docx.Document()



# adding title for the recipe

tacoWord.add_paragraph('Random Taco Cookbook', 'Title')



# starting new blank page

tacoWord.add_picture('modified_taco.jpg')



# To print this under picture

tacoWord.add_paragraph("Tai's Captures")



tacoWord.add_paragraph('first last')



# break pages in all five recipe components

tacoWord.add_page_break()



# URL source for picture

tacoWord.add_paragraph('https://taco-1150.herokuapp.com/random/?full_taco=true')





# Five components of the recipe and loops



for item in range(3):  # to pull the five recipes



    # for printing the seasoning recipes

    tacoWord.add_paragraph(tacoBook[item]["seasoning"]["name"], 'heading 1')



    # for seasoning recipes

    tacoWord.add_paragraph(tacoBook[item]["seasoning"]["recipe"])



    # for condiment recipes

    tacoWord.add_paragraph(tacoBook[item]["condiment"]["recipe"])



    # for mixin recipes

    tacoWord.add_paragraph(tacoBook[item]["mixin"]["recipe"])



    # for base layer recipes

    tacoWord.add_paragraph(tacoBook[item]["base_layer"]["recipe"])



    # end cycle

    tacoWord.add_page_break()



# save word document

tacoWord.save('recipe.docx')
